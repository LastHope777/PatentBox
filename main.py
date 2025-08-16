import os

from selenium import webdriver
import time
from docx import Document
import sys
from PyQt5 import QtWidgets, QtGui
from patent_ptoject_design_main_menu import Ui_MainWindow
from PyQt5.QtWidgets import QMessageBox
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from PyQt5.QtWidgets import QFileDialog
from patent_ptoject_design_number import Ui_SecondWindow


PathToDriver = ""
class MyApp(QtWidgets.QMainWindow):

    # Главный экрана
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.second_window = None
        self.ui.pushButton_start.clicked.connect(self.open_second_window)

    # Переход на второй экран
    def open_second_window(self):
        self.second_window = WorkWithWord()
        self.second_window.show()
        self.close()


class WorkWithWord(QtWidgets.QMainWindow):

    # Второй экран
    def __init__(self):
        super().__init__()
        self.ui = Ui_SecondWindow()
        self.ui.setupUi(self)

        # Привязка кнопок выбора файлов
        self.ui.input_btn.clicked.connect(self.open_file_dialog_input)
        self.ui.adapter_btn.clicked.connect(self.open_file_dialog_driver)

        # Логика чекбоксов
        self.ui.cb7.stateChanged.connect(self.toggle_checkboxes)

        # Кнопка старта
        self.ui.start_btn.clicked.connect(self.WorkWithWord_run)

        # Кнопка "Назад"
        self.ui.back_btn.clicked.connect(self.go_back)

    def go_back(self):
        self.first_window = MyApp()
        self.first_window.show()
        self.close()

    def show_success_message(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle("Успех")
        msg.setText("Данные успешно сохранены")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

    def toggle_checkboxes(self):
        is_checked = self.ui.cb7.isChecked()
        for cb in [self.ui.cb1, self.ui.cb2, self.ui.cb3,
                   self.ui.cb4, self.ui.cb5, self.ui.cb6]:
            cb.setChecked(is_checked)

    def open_file_dialog_input(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл",
            "",
            "Документы (*.docx)"
        )
        if file_path:
            self.ui.input_path_edit.setText(file_path)

    def open_file_dialog_driver(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл",
            "",
            "Программы (*.exe)"
        )
        if file_path:
            self.ui.adapter_path_edit.setText(file_path)

    def WorkWithWord_run(self):
        path_to_input = self.ui.input_path_edit.text()
        path_to_driver = self.ui.adapter_path_edit.text()

        selected_checkboxes = []
        if self.ui.cb1.isChecked():
            selected_checkboxes.append('1')
        if self.ui.cb2.isChecked():
            selected_checkboxes.append('2')
        if self.ui.cb3.isChecked():
            selected_checkboxes.append('3')
        if self.ui.cb4.isChecked():
            selected_checkboxes.append('4')
        if self.ui.cb5.isChecked():
            selected_checkboxes.append('5')
        if self.ui.cb6.isChecked():
            selected_checkboxes.append('6')
        if self.ui.cb7.isChecked():
            selected_checkboxes.append('7')

        WordMode(path_to_input, selected_checkboxes, path_to_driver)
        self.show_success_message()


class Patent:

    def __init__(self, number, name, authors, mpk, date):
        self.number = number
        self.name = name
        self.authors = authors
        self.MPK = mpk
        self.date = date

def WordMode(path_to_document, selected_checkboxes, path_to_driver):
    result = []
    document = Document(path_to_document)
    for table in document.tables:
        if (table.cell(0, 0).text == 'ФИПС'
                or table.cell(0, 0).text == 'Фипс'
                or table.cell(0, 0).text == 'фипс'):

            number_search = []
            arr_number = []
            count = 0
            for rows in table.rows[1:]:
                cell_text = rows.cells[0].text
                words = cell_text.split(" ")
                if len(words) > 0:
                    arr_number.append(words)
                    number_search.append(words[0])
            document_result = Document("result.docx")
            first_empty_cell = 2
            table_fips = document_result.add_table(rows=2, cols=6)
            table_fips.cell(0, 0).text = 'Объект исследования, его составные части'
            table_fips.cell(0,
                           1).text = ('Вид и номер охранного документа с двухбуквенным кодом страны/режим действия ('
                                      'действует/не действует/может быть восстановлен)')
            table_fips.cell(0,
                           2).text = 'Название объекта интеллектуальной собственности'
            table_fips.cell(0,
                           3).text = 'Заявитель (правообладатель)'
            table_fips.cell(0,
                           4).text = 'Классификационные рубрики'
            table_fips.cell(0,
                           5).text = 'Дата приоритета/ дата публикации'
            table_fips.cell(1, 0).text = '1'
            table_fips.cell(1,
                           1).text = '2'
            table_fips.cell(1,
                           2).text = '3'
            table_fips.cell(1,
                           3).text = '4'
            table_fips.cell(1,
                           4).text = '5'
            table_fips.cell(1,
                           5).text = '6'
            for request in number_search:
                url = 'https://www.fips.ru/iiss/'
                service = Service(executable_path=path_to_driver)
                driver = webdriver.Chrome(service=service)
                driver.get(url)
                time.sleep(2)
                patent_rus = driver.find_element(By.XPATH,
                                                 '//*[@id="db-selection-form:j_idt74"]/div')
                patent_rus.click()
                if '7' in selected_checkboxes:
                    patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:j_idt101"]')
                    patent_rus.click()
                else:
                    if '1' in selected_checkboxes:
                        patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:dbsGrid1:0'
                                                                   ':dbsGrid1checkbox"]')
                        patent_rus.click()
                    if '2' in selected_checkboxes:
                        patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:dbsGrid1:1'
                                                                   ':dbsGrid1checkbox"]')
                        patent_rus.click()
                    if '3' in selected_checkboxes:
                        patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:dbsGrid1:2'
                                                                   ':dbsGrid1checkbox"]')
                        patent_rus.click()
                    if '4' in selected_checkboxes:
                        patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:dbsGrid1:3'
                                                                   ':dbsGrid1checkbox"]')
                        patent_rus.click()
                    if '5' in selected_checkboxes:
                        patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:dbsGrid1:4'
                                                                   ':dbsGrid1checkbox"]')
                        patent_rus.click()
                    if '6' in selected_checkboxes:
                        patent_rus = driver.find_element(By.XPATH, '//*[@id="db-selection-form:dbsGrid1:5'
                                                                   ':dbsGrid1checkbox"]')
                        patent_rus.click()

                time.sleep(4)
                button_search = driver.find_element(By.XPATH, '//*[@id="db-selection-form:button-set1"]/div[1]/input')
                button_search.click()
                time.sleep(4)
                main_area_request = driver.find_element(By.XPATH, '//*[@id="fields:1:j_idt109"]')
                main_area_request.click()
                main_area_request.send_keys(request)
                time.sleep(1)
                button_search = driver.find_element(By.XPATH,
                                                    '// *[ @ id = "searchForm"] / div[1] / div[1] / div[3] / div[1] / input')
                button_search.click()
                time.sleep(2)
                result1 = driver.find_element(By.XPATH,
                                              '/html/body/div[3]/div/div/div[1]/div[2]/div/form/div[3]/div/div/a/div[1]')
                result1.click()


                try:
                    validity = driver.find_element(By.XPATH, '//*[@id="StatusR"]').text
                except:
                    try:
                        validity = driver.find_element(By.XPATH, '//*[@id="StatusRAP"]').text
                    except:
                        validity = "Не удалось найти информацию о режиме действия патента"
                number = (driver.find_element(By.XPATH, '//*[@id="top2"]').text + " "
                          + driver.find_element(By.XPATH, '//*[@id="top4"]').text + " "
                          + driver.find_element(By.XPATH, '//*[@id="top6"]').text + "\n" + validity)
                count += 1
                name = driver.find_element(By.XPATH,
                                           '/html/body/div[3]/div/div/div[1]/div[2]/form/div/div/div[2]/div/div/p/b').text
                authors = driver.find_element(By.XPATH, '// *[ @ id = "bibl"] / p[1] / b').text
                mpk = driver.find_element(By.XPATH,
                                          '// *[ @ id = "mainDoc"] / table[1] / tbody / tr / td[2] / table / tbody / '
                                          'tr[2] / td[1] / div / ul / li / a / span').text
                date = driver.find_element(By.XPATH, '// *[ @ id = "bib"] / tbody / tr / td[1] / p[2] / b').text
                time.sleep(1)
                result.append(Patent(number, name, authors, mpk, date))

                table_fips.add_row()
                for _ in table_fips.rows[1:]:
                    table_fips.cell(first_empty_cell, 1).text = number
                    table_fips.cell(first_empty_cell, 2).text = name
                    table_fips.cell(first_empty_cell, 3).text = authors
                    table_fips.cell(first_empty_cell, 4).text = mpk
                    table_fips.cell(first_empty_cell, 5).text = date
                first_empty_cell += 1
                document_result.save("result.docx")

                driver.quit()

        if table.cell(0, 0).text == 'Платформа' or table.cell(0, 0).text == 'платформа':
            url = "https://openstat.rospatent.gov.ru/patents"
            number_search = []
            document_result = Document("result.docx")
            first_empty_cell = 2
            table_platform = document_result.add_table(rows=2, cols=6)
            table_platform.cell(0, 0).text = 'Объект исследования, его составные части'
            table_platform.cell(0,
                                1).text = ('Вид и номер охранного документа с двухбуквенным кодом страны/режим '
                                           'действия (действует/не действует/может быть восстановлен)')
            table_platform.cell(0,
                                2).text = 'Название объекта интеллектуальной собственности'
            table_platform.cell(0,
                                3).text = 'Заявитель (правообладатель)'
            table_platform.cell(0,
                                4).text = 'Классификационные рубрики'
            table_platform.cell(0,
                                5).text = 'Дата приоритета/ дата публикации'
            table_platform.cell(1, 0).text = '1'
            table_platform.cell(1,
                                1).text = '2'
            table_platform.cell(1,
                                2).text = '3'
            table_platform.cell(1,
                                3).text = '4'
            table_platform.cell(1,
                                4).text = '5'
            table_platform.cell(1,
                                5).text = '6'
            for rows in table.rows[1:]:
                cell_text = rows.cells[0].text
                number_search.append(cell_text)
            service = Service(executable_path=path_to_driver)
            driver = webdriver.Chrome(service=service)
            for request in number_search:
                driver.get(url)
                time.sleep(2)
                search_input = driver.find_element(By.XPATH,
                                                   '//*[@id="layout"]/div[3]/div/div[3]/div[1]/div[1]/div[1]/input')
                search_input.clear()
                search_input.send_keys(request)
                button_search_platform = driver.find_element(By.XPATH,
                                                              '//*[@id="layout"]/div[3]/div/div[2]/div/div/button')
                button_search_platform.click()
                time.sleep(2)
                result_patent = driver.find_element(By.XPATH,
                                                    '/html/body/div/div/div[3]/div/div[4]/div[4]/div/div['
                                                    '2]/div/div/div/ul/li[1]/div[2]/div/div[1]/div/div')
                result_patent.click()
                time.sleep(2)
                country = driver.find_element(By.XPATH, '//*[@id="doc-biblio"]/div[2]/div[1]/div[1]/div[2]')
                number1 = driver.find_element(By.XPATH, '//*[@id="doc-biblio"]/div[2]/div[1]/div[2]/div[2]')
                number2 = driver.find_element(By.XPATH, '//*[@id="doc-biblio"]/div[2]/div[1]/div[3]/div[2]')
                number = country.text + " " + number1.text + " " + number2.text
                name = driver.find_element(By.XPATH,
                                           '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div[3]/div/div['
                                           '1]/div/div[1]/div[1]/h1').text
                try:
                    # Патентообладатель
                    try:
                        authors = driver.find_element(By.XPATH,
                                                      '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                      '3]/div/div[1]/div/div[1]/div[2]/div[6]/div[3]/ul').text
                    except:
                        authors = driver.find_element(By.XPATH,
                                                      '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                      '3]/div/div[1]/div/div[1]/div[2]/div[5]/div[3]/ul').text
                except:
                    try:
                        # Заявитель
                        try:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                          '3]/div/div[1]/div/div[1]/div[2]/div[4]/div[3]/ul').text
                        except:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                          '3]/div/div[1]/div/div[1]/div[2]/div[3]/div[3]/ul').text
                    except:
                        # Автор
                        try:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                          '3]/div/div[1]/div/div[1]/div[2]/div[3]/div[3]/ul').text
                        except:
                            try:
                                authors = driver.find_element(By.XPATH,
                                                              '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                              '3]/div/div[1]/div/div[1]/div[2]/div[4]/div[3]/ul').text
                            except:
                                authors = driver.find_element(By.XPATH,
                                                              '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div['
                                                              '3]/div/div[1]/div/div[1]/div[2]/div[5]/div[3]/ul').text
                mpk_text = driver.find_element(By.XPATH,
                                          '//*[@id="doc-biblio"]/div[2]/div[3]/div[2]').text
                if mpk_text == "МПК":
                    mpk = driver.find_element(By.XPATH,
                                            '//*[@id="doc-biblio"]/div[2]/div[3]/div[3]').text
                else:
                    mpk = driver.find_element(By.XPATH,
                                            '//*[@id="doc-biblio"]/div[2]/div[4]/div[3]').text
                try:
                    date = driver.find_element(By.XPATH,
                                               '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div[3]/div/div['
                                               '1]/div/div[1]/div[1]/div[5]/div[4]').text
                except:
                    date = driver.find_element(By.XPATH,
                                               '/html/body/div/div/div[3]/div/div/div[2]/div[1]/div[3]/div/div['
                                               '1]/div/div[1]/div[1]/div[4]/div[4]').text
                result.append(Patent(number, name, authors, mpk, date))
                table_platform.add_row()
                for _ in table_platform.rows[1:]:
                    table_platform.cell(first_empty_cell, 1).text = number
                    table_platform.cell(first_empty_cell, 2).text = name
                    table_platform.cell(first_empty_cell, 3).text = authors
                    table_platform.cell(first_empty_cell, 4).text = mpk
                    table_platform.cell(first_empty_cell, 5).text = date
                first_empty_cell += 1
                document_result.save("result.docx")

            driver.quit()

        if (table.cell(0, 0).text == 'WIPO'
                or table.cell(0, 0).text == 'wipo'
                or table.cell(0, 0).text == 'Wipo'):

            number_search = []
            document_result = Document("result.docx")
            first_empty_cell = 2
            table_wipo = document_result.add_table(rows=2, cols=6)

            table_wipo.cell(0, 0).text = 'Объект исследования, его составные части'
            table_wipo.cell(0,
                           1).text = ('Вид и номер охранного документа с двухбуквенным кодом страны/режим действия ('
                                      'действует/не действует/может быть восстановлен)')
            table_wipo.cell(0,
                           2).text = 'Название объекта интеллектуальной собственности'
            table_wipo.cell(0,
                           3).text = 'Заявитель (правообладатель)'
            table_wipo.cell(0,
                           4).text = 'Классификационные рубрики'
            table_wipo.cell(0,
                           5).text = 'Дата приоритета/ дата публикации'
            table_wipo.cell(1, 0).text = '1'
            table_wipo.cell(1,
                           1).text = '2'
            table_wipo.cell(1,
                           2).text = '3'
            table_wipo.cell(1,
                           3).text = '4'
            table_wipo.cell(1,
                           4).text = '5'
            table_wipo.cell(1,
                           5).text = '6'

            for rows in table.rows[1:]:
                cell_text = rows.cells[0].text
                number_search.append(cell_text)
            for request in number_search:

                service = Service(executable_path=path_to_driver)
                driver = webdriver.Chrome(service=service)
                url = "https://patentscope.wipo.int/search/ru/search.jsf"
                driver.get(url)
                time.sleep(5)
                field_search = driver.find_element(By.XPATH, '//*[@id="simpleSearchForm:fpSearch:input"]')
                field_search.send_keys(request)
                button_search_wipo = driver.find_element(By.XPATH,
                                                         '/html/body/div[2]/div[5]/div/div[2]/form/div/div[1]/div['
                                                         '2]/div/div/div[1]/div[2]/button')
                button_search_wipo.click()
                time.sleep(5)

                try:
                    name = driver.find_element(By.XPATH,
                                               '//*[@id="headerForm:headerFormContent"]/h1/div/div[1]').text
                    suffix1 = name.index(" - ") + 3
                    str3 = name[suffix1:]
                    number = request
                    name = str3

                    try:
                        authors = driver.find_element(By.XPATH,
                                                      '/html/body/div[2]/div[5]/div/div[1]/div['
                                                      '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                      '10]/span[2]/span').text
                    except:
                        try:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div[2]/div[5]/div/div[1]/div['
                                                          '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                          '9]/span[2]/span').text
                        except:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div[2]/div[5]/div/div[1]/div['
                                                          '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                          '11]/span[2]/span').text

                    try:
                        mpk = driver.find_element(By.XPATH,
                                                  '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                                  '1]/div/div/div[2]/div/div[1]/div[9]/span[2]/div/div[1]').text
                    except:
                        try:
                            mpk = driver.find_element(By.XPATH,
                                                      '/html/body/div[2]/div[5]/div/div[1]/div['
                                                      '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                      '5]/span[2]/div/div[1]').text
                        except:
                            mpk = driver.find_element(By.XPATH,
                                                      '/html/body/div[2]/div[5]/div/div[1]/div['
                                                      '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                      '7]/span[2]/div/div[1]/div').text

                    date = driver.find_element(By.XPATH,
                                               '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                               '1]/div/div/div[2]/div/div[1]/div[5]/span[2]').text
                except:
                    result2 = driver.find_element(By.XPATH,
                                                  '//*[@id="resultListForm:resultTable:0:patentResult"]/div[1]/div['
                                                  '1]/a/span')

                    result2.click()
                    time.sleep(4)
                    number1 = driver.find_element(By.XPATH,
                                                  '//*[@id="headerForm:headerFormContent"]/h1/div/div[1]').text
                    suffix1 = number1.index(". ") + 2
                    str1 = number1[suffix1:]
                    suffix2 = number1.index(" - ") - 3
                    str2 = str1[:suffix2]
                    try:
                        temp = driver.find_element(By.XPATH,
                                                   '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                                   '1]/div/div/div[2]/div/div[1]/div[6]/span[1]/span').text
                        if temp == "Вид публикации":
                            number2 = driver.find_element(By.XPATH,
                                                          '/html/body/div[2]/div[5]/div/div[1]/div['
                                                          '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                          '8]/span[2]').text
                        else:
                            number2 = driver.find_element(By.XPATH,
                                                          '/html/body/div[2]/div[5]/div/div[1]/div['
                                                          '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                          '6]/span[2]').text
                    except:
                        number2 = ""
                    name = driver.find_element(By.XPATH,
                                               '//*[@id="headerForm:headerFormContent"]/h1/div/div[1]').text
                    suffix1 = name.index(" - ") + 3
                    str3 = name[suffix1:]
                    number = str2 + " " + number2
                    name = str3
                    try:
                        authors = driver.find_element(By.XPATH,
                                                      '/html/body/div[2]/div[5]/div/div[1]/div['
                                                      '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                      '10]/span[2]/span').text
                    except:
                        try:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div[2]/div[5]/div/div[1]/div['
                                                          '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                          '9]/span[2]/span').text
                        except:
                            authors = driver.find_element(By.XPATH,
                                                          '/html/body/div[2]/div[5]/div/div[1]/div['
                                                          '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                          '11]/span[2]/span').text

                    try:
                        mpk = driver.find_element(By.XPATH,
                                                  '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                                  '1]/div/div/div[2]/div/div[1]/div[9]/span[2]/div/div[1]').text
                    except:
                        try:
                            mpk = driver.find_element(By.XPATH,
                                                      '/html/body/div[2]/div[5]/div/div[1]/div['
                                                      '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                      '5]/span[2]/div/div[1]').text
                        except:
                            mpk = driver.find_element(By.XPATH,
                                                      '/html/body/div[2]/div[5]/div/div[1]/div['
                                                      '2]/form/div/div/div/div[1]/div/div/div[2]/div/div[1]/div['
                                                      '7]/span[2]/div/div[1]/div').text
                    date = driver.find_element(By.XPATH,
                                               '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                               '1]/div/div/div[2]/div/div[1]/div[5]/span[2]').text
                result.append(Patent(number, name, authors, mpk, date))

                table_wipo.add_row()
                if mpk == date:
                    authors = driver.find_element(By.XPATH,
                                              '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                              '1]/div/div/div[2]/div/div[1]/div[7]/span[2]/span/ul').text
                    mpk = driver.find_element(By.XPATH,
                                              '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                              '1]/div/div/div[2]/div/div[1]/div[5]/span[2]/div/div[1]').text
                    date = driver.find_element(By.XPATH,
                                               '/html/body/div[2]/div[5]/div/div[1]/div[2]/form/div/div/div/div['
                                               '1]/div/div/div[2]/div/div[1]/div[2]/span[2]').text
                for _ in table_wipo.rows[1:]:
                    table_wipo.cell(first_empty_cell, 1).text = number
                    table_wipo.cell(first_empty_cell, 2).text = name
                    table_wipo.cell(first_empty_cell, 3).text = authors
                    table_wipo.cell(first_empty_cell, 4).text = mpk
                    table_wipo.cell(first_empty_cell, 5).text = date

                first_empty_cell += 1
                document_result.save("result.docx")

                driver.quit()

    return result


if __name__ == "__main__":

    app = QtWidgets.QApplication(sys.argv)
    icon_path = os.path.join(os.getcwd(), "logo.png")
    app.setWindowIcon(QtGui.QIcon(icon_path))
    window = MyApp()
    window.show()
    sys.exit(app.exec_())
